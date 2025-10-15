from dotenv import load_dotenv
load_dotenv()

import os
import tempfile
import vertexai
from typing import Optional

from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.templating import Jinja2Templates
from fastapi.responses import FileResponse
from starlette.background import BackgroundTask

from vertexai.generative_models import GenerativeModel, Part
from html4docx import HtmlToDocx
from docx import Document
from docx.shared import Inches

# --- Configuration ---
PROJECT_ID = os.environ.get("GOOGLE_CLOUD_PROJECT")
REGION = os.environ.get("GOOGLE_CLOUD_REGION")

if not PROJECT_ID or not REGION:
    raise ValueError("Missing required environment variables: GOOGLE_CLOUD_PROJECT and GOOGLE_CLOUD_REGION")

vertexai.init(project=PROJECT_ID, location=REGION)

# --- Helper Functions ---
def load_prompt(file_name: str) -> str:
    """Loads a prompt template from the prompts directory."""
    try:
        with open(f"prompts/{file_name}", "r") as f:
            return f.read()
    except FileNotFoundError:
        return f"Error: Prompt file '{file_name}' not found."

# --- Constants ---
GEMINI_MODELS = {
    "gemini-2.5-flash": {
        "name": "Gemini 2.5 Flash",
        "pricing_input": "$0.075 / 1M tokens",
        "pricing_output": "$0.30 / 1M tokens"
    },
    "gemini-2.5-pro": {
        "name": "Gemini 2.5 Pro",
        "pricing_input": "$7.00 / 1M tokens",
        "pricing_output": "$21.00 / 1M tokens"
    }
}

BASE_PROMPT = load_prompt("base_prompt.txt")
CORRECTION_PROMPT = load_prompt("correction_prompt.txt")


from pydantic import BaseModel


import re

# --- Data Models ---
class PromptOptions(BaseModel):
    signature_handling: str
    seal_handling: str
    currency_format: str
    numbers_format: str
    general_instructions: str


class SavePromptOptions(BaseModel):
    filename: str
    content: str


# --- FastAPI App Setup ---
app = FastAPI()
templates = Jinja2Templates(directory="app/templates")


@app.get("/")
def read_root(request: Request):
    """Serves the initial HTML page and provides model information."""
    return templates.TemplateResponse("index.html", {"request": request, "models": GEMINI_MODELS})


@app.post("/preview-prompt")
async def preview_prompt(options: PromptOptions):
    """Generates and returns a formatted prompt based on user options."""
    prompt = BASE_PROMPT.format(
        signature_handling=options.signature_handling,
        seal_handling=options.seal_handling,
        currency_format=options.currency_format,
        numbers_format=options.numbers_format,
        general_instructions=options.general_instructions
    )
    return {"prompt": prompt}


@app.get("/list-prompts")
async def list_prompts():
    """Lists all available prompt files."""
    try:
        files = [f for f in os.listdir("prompts") if f.endswith(".txt")]
        return {"prompts": files}
    except FileNotFoundError:
        return {"prompts": []}


@app.get("/load-prompt/{filename}")
async def load_prompt_file(filename: str):
    """Loads a specific prompt file's content."""
    # Basic security: prevent directory traversal
    if ".." in filename or "/" in filename:
        return {"error": "Invalid filename"}
    try:
        with open(f"prompts/{filename}", "r") as f:
            return {"content": f.read()}
    except FileNotFoundError:
        return {"error": "File not found"}


@app.post("/save-prompt")
async def save_prompt(options: SavePromptOptions):
    """Saves content to a new prompt file."""
    # Sanitize filename: allow letters, numbers, underscores, hyphens
    safe_filename = re.sub(r'[^a-zA-Z0-9_\-]', '', options.filename)
    if not safe_filename:
        return {"error": "Invalid filename provided."}
    
    safe_filename += ".txt"
    
    try:
        with open(f"prompts/{safe_filename}", "w") as f:
            f.write(options.content)
        return {"success": True, "filename": safe_filename}
    except Exception as e:
        return {"error": str(e)}


@app.post("/upload")
async def upload_file(
    file: UploadFile = File(...),
    reference_docx: Optional[UploadFile] = File(None),
    model_choice: str = Form("gemini-2.5-flash"),
    margin_top: Optional[float] = Form(1.0),
    margin_bottom: Optional[float] = Form(1.0),
    margin_left: Optional[float] = Form(1.0),
    margin_right: Optional[float] = Form(1.0),
    custom_prompt: str = Form(...),
    needs_correction: bool = Form(False),
    correction_instructions: str = Form("None"),
):
    """Handles file upload, AI processing, DOCX conversion, and advanced formatting."""

    # --- Initial AI Conversion ---
    image_part = Part.from_data(mime_type=file.content_type, data=await file.read())
    
    prompt = custom_prompt
    
    model = GenerativeModel(model_choice)
    response = await model.generate_content_async([image_part, prompt])
    html_content = response.text

    # --- Optional Correction Run ---
    if needs_correction and correction_instructions != "None":
        correction_prompt = CORRECTION_PROMPT.format(
            correction_instructions=correction_instructions,
            original_html=html_content
        )
        response = await model.generate_content_async(correction_prompt)
        html_content = response.text

    # --- DOCX Generation and Formatting ---
    
    # Clean up potential markdown fences from the AI response
    if "```html" in html_content:
        html_content = html_content.split("```html")[1].split("```")[0]
    elif "```" in html_content:
        html_content = html_content.split("```")[1].split("```")[0]

    document = Document()

    # --- Apply Margins ---
    if reference_docx and reference_docx.filename != '':
        try:
            ref_doc = Document(reference_docx.file)
            if ref_doc.sections:
                section = document.sections[0]
                ref_section = ref_doc.sections[0]
                section.top_margin = ref_section.top_margin
                section.bottom_margin = ref_section.bottom_margin
                section.left_margin = ref_section.left_margin
                section.right_margin = ref_section.right_margin
        except Exception as e:
            print(f"Could not apply reference doc margins: {e}")
            # Fallback to default or provided margins if reference doc fails
            section = document.sections[0]
            section.top_margin = Inches(margin_top)
            section.bottom_margin = Inches(margin_bottom)
            section.left_margin = Inches(margin_left)
            section.right_margin = Inches(margin_right)
    else:
        section = document.sections[0]
        section.top_margin = Inches(margin_top)
        section.bottom_margin = Inches(margin_bottom)
        section.left_margin = Inches(margin_left)
        section.right_margin = Inches(margin_right)

    # --- Convert HTML to DOCX ---
    parser = HtmlToDocx()

    # Set default font to Arial
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'

    parser.add_html_to_document(html_content, document)

    # Remove heading styles to flatten the document structure
    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            paragraph.style = document.styles['Normal']

    # --- Post-process Tables ---
    for table in document.tables:
        table.autofit = True
        table.allow_autofit = True

    # --- Save and Return File ---
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        document.save(temp_docx)
        temp_docx_path = temp_docx.name

    return FileResponse(
        path=temp_docx_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=f"{os.path.splitext(file.filename)[0]}_converted.docx",
        background=BackgroundTask(lambda: os.remove(temp_docx_path))
    )
